from docx import Document
from docx.shared import Pt, RGBColor #字体大小和颜色
from docx.enum.text import WD_ALIGN_PARAGRAPH #调整段落居中
import datetime
guest_name = 'Taylor Swift'
doc = Document()
header = doc.add_paragraph()
header_run = header.add_run('INVITATION')
header_run.font.bold = True
header_run.font.size = Pt(30)
header.alignment = WD_ALIGN_PARAGRAPH.CENTER

title = doc.add_paragraph()
title_run = title.add_run(guest_name.title()+' :')
title_run.font.size = Pt(20)
title_run.font.name = 'Segoe Script'
content = doc.add_paragraph()
content_run = content.add_run("Welcome to my party that will begin at 17:00 on April 1st: "\
"It will take place at No.94 Nebulas, Virgo. I would really like you could"\
"come and enjoy the time with us." #字符串太长，使用\来连接多行
)
content_run.font.size = Pt(12)
content_run.font.name = 'Segoe Script'
content.paragraph_format.line_spacing = 2.0
signature = doc.add_paragraph()
signature_run = signature.add_run('Nisha Qiao\n')
signature_run.font.name = 'Segoe  Script'
signature_run.font.size = Pt(18)
now_time = datetime.datetime.now().strftime('%Y-%m-%d')
sign_time_run = signature.add_run(now_time)
sign_time_run.font.size = Pt(12)
sign_time_run.font.name = 'Segoe  Script'
signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
doc.save('invitation_to_'+guest_name+'.docx')
